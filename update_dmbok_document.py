from copy import deepcopy
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "M-AI Master技术说明文档V4.3.docx"
TARGET = ROOT / "M-AI Master技术说明文档V5.1_企业Agent增强版_最新版.docx"


def set_cell_shading(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def style_table(table, widths=None):
    table.style = "Table Grid"
    table.autofit = True
    for index, cell in enumerate(table.rows[0].cells):
        set_cell_shading(cell, "1769AA")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)
    for row in table.rows[1:]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                if index < len(row.cells):
                    row.cells[index].width = width


def add_paragraph(doc, elements, text="", style=None, page_break=False):
    paragraph = doc.add_paragraph(style=style)
    if text:
        paragraph.add_run(text)
    if page_break:
        paragraph.add_run().add_break(WD_BREAK.PAGE)
    elements.append(paragraph._p)
    return paragraph


def add_table(doc, elements, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
    style_table(table)
    elements.append(table._element)
    return table


def replace_text(doc, old, new):
    for paragraph in doc.paragraphs:
        if old in paragraph.text:
            if len(paragraph.runs) == 1:
                paragraph.runs[0].text = paragraph.runs[0].text.replace(old, new)
            else:
                style = paragraph.style
                paragraph.text = paragraph.text.replace(old, new)
                paragraph.style = style
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if old in paragraph.text:
                        style = paragraph.style
                        paragraph.text = paragraph.text.replace(old, new)
                        paragraph.style = style


def set_paragraph(doc, startswith, text):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(startswith):
            style = paragraph.style
            paragraph.text = text
            paragraph.style = style
            return True
    return False


def remove_table_rows_containing(doc, terms):
    for table in doc.tables:
        for row in list(table.rows):
            text = " | ".join(cell.text for cell in row.cells)
            if any(term.lower() in text.lower() for term in terms):
                table._tbl.remove(row._tr)


def remove_tables_containing(doc, terms):
    for table in list(doc.tables):
        text = " | ".join(cell.text for row in table.rows for cell in row.cells)
        if any(term.lower() in text.lower() for term in terms):
            table._tbl.getparent().remove(table._tbl)


def build_document():
    doc = Document(str(SOURCE))
    doc.core_properties.title = "M-AI Master 技术说明文档 V5.1 企业 Agent 增强版"
    doc.core_properties.subject = "统一主数据 Agent、企业 RAG、可信血缘与 DMBOK2 工程控制映射"
    doc.core_properties.comments = "2026-08-08 基于当前可运行代码、可视验证和自动化测试证据更新"
    doc.core_properties.modified = datetime.now(timezone.utc)

    replace_text(doc, "技术说明文档 V4.3", "技术说明文档 V5.1")
    replace_text(doc, "M-AI Master 4.3", "M-AI Master 5.1")
    replace_text(doc, "比赛与企业交付增强版", "统一 Agent 与企业可信治理增强版")
    replace_text(doc, "24/24", "40/40")
    replace_text(
        doc,
        "16 项 API 回归 + 8 项故障与降级测试",
        "21 项 API 回归 + 9 项故障与降级测试 + 10 项企业治理测试",
    )
    replace_text(doc, "4.3、就绪", "5.1、就绪")
    replace_text(doc, "version=4.3", "version=5.1")
    replace_text(
        doc,
        "16/16 API 回归和 8/8 故障测试",
        "21/21 API 回归、9/9 故障测试和 10/10 企业治理测试",
    )
    replace_text(
        doc,
        "多编码接入；模拟 OCR；SAP/EAM/MES/WMS 载荷与日志",
        "多编码接入；PaddleOCR → Qwen-VL → 规则三级识别；SAP/EAM/MES/WMS 载荷与日志",
    )
    replace_text(doc, "24 项自动化测试", "40 项自动化测试")
    replace_text(
        doc,
        "16 项 API 回归 + 8 项故障测试",
        "21 项 API 回归 + 9 项故障与降级测试 + 10 项企业治理测试",
    )
    replace_text(
        doc,
        "50 条、27 组人工归并真值集 + 两套可重复治理基线",
        "业务签字验证集尚待建立；当前以两套可重复治理基线验证流程稳定性",
    )
    replace_text(
        doc,
        "样本规模仍小，Pairwise F1 尚有提升空间，未形成跨数据域与长期业务工时证据",
        "尚未形成业务签字准确率、跨数据域长期工时与收益证据",
    )
    replace_text(
        doc,
        "扩展分层金标准集，按品类/工厂报告 Precision、Recall、F1、置信区间与真实节省工时",
        "由业务方建立独立签字验证集，按品类/工厂报告 Precision、Recall、F1、置信区间与真实节省工时",
    )
    replace_text(
        doc,
        "POST /api/agent/plan；GET /api/workflow/{batch_id}",
        "POST /api/agent/query；POST /api/agent/transcribe；POST /api/agent/plan；GET /api/workflow/{batch_id}",
    )
    replace_text(
        doc,
        "计划、八步状态、进度、closed_loop",
        "意图、实体、计划、八步状态、进度、closed_loop",
    )
    replace_text(
        doc,
        "条件、语义元数据、结果和无精确匹配建议",
        "条件、语义元数据、主数据结果、knowledge_references 标准依据和模糊候选建议",
    )
    replace_text(
        doc,
        "4.3、就绪、存储、模型、OCR、工厂、部署和认证状态",
        "5.1、就绪、存储、模型、OCR、工厂、部署和认证状态",
    )
    remove_tables_containing(doc, ("代码位置", "evaluation/ground_truth_50.csv"))
    replace_text(
        doc,
        "审计能力是单节点 SHA-256 防篡改哈希链，不是公链或多节点联盟链",
        "业务审计采用批次 SHA-256 指纹链，安全事件采用 HMAC 签名链；两者都不是公链或多节点联盟链",
    )
    replace_text(
        doc,
        "互联网部署必须增加 HTTPS、SSO/RBAC 与集中密钥管理",
        "互联网部署仍需增加 HTTPS、企业 SSO、集中密钥管理和审计外送",
    )
    replace_text(
        doc,
        "基础认证、Trace ID、轮转日志",
        "四账号会话、RBAC、工厂域、CSRF、HMAC 审计、Trace ID 与轮转日志",
    )
    replace_text(
        doc,
        "尚未接入 SSO/RBAC、HTTPS 与集中日志",
        "已实现本地企业认证与权限控制；尚未接入企业 SSO、HTTPS 与集中日志",
    )
    remove_table_rows_containing(
        doc,
        ("/api/evaluation/", "evaluation/ground_truth", "人工真值评测", "人工真值集与阈值校准"),
    )

    set_paragraph(
        doc,
        "本文档是 M-AI Master 5.1",
        "本文档是 M-AI Master 5.1 当前可运行系统的统一 Agent 与企业可信治理增强版技术说明，依据 Flask 后端、单页前端、SQLite 权威数据库、文本/语音/图片统一任务入口、版本化 RAG 标准知识库、关系图与桑基流双视图、企业身份与工厂范围控制、数据质量问题闭环、签名审计链和 40 项自动化测试更新。文档区分工程控制、业务验收与正式合规测评：工程证据可由接口和测试复核，业务准确率与投资回报仍需真实岗位基线和业务方签字，等保、法律合规及 DAMA 能力认证仍需部署单位另行评估。",
    )
    set_paragraph(
        doc,
        "运行结果：test_api.py",
        "运行结果：test_api.py 为 21/21，test_chaos.py 为 9/9，test_enterprise.py 为 10/10，总计 40/40，全部 OK。测试覆盖动态数据接入、统一 Agent 意图路由、单任务分发、RAG 联合检索、语义降级、真实 OCR 适配、八步闭环、四角色权限、工厂越权阻断、CSRF、法律保全、审计篡改检测、DMBOK 控制目录和质量问题闭环。",
    )

    stale_phrases = (
        "50 条、27 组人工真值集",
        "50条人工真值集",
        "50 条人工真值集",
    )
    for paragraph in doc.paragraphs:
        if any(phrase in paragraph.text for phrase in stale_phrases):
            paragraph.text = (
                "当前版本不把结构完整性、候选压缩率或算法代理分冒充业务准确率。"
                "Precision、Recall 和经济收益必须基于业务专家签字的独立验证集与真实岗位计时；"
                "该验证属于企业试点验收输入，不在演示数据中预设结论。"
            )

    toc = {
        "19. 核心实现代码": "19. DMBOK2 数据管理审核与整改",
        "20. 测试证据与代码定位": "20. 核心实现代码",
    }
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in toc:
            paragraph.text = toc[text]
    directory = next((p for p in doc.paragraphs if p.text.strip() == "目录"), None)
    if directory is not None:
        siblings = list(directory._p.getparent())
        position = siblings.index(directory._p)
        toc_paragraphs = []
        for sibling in siblings[position + 1 :]:
            if sibling.tag == qn("w:p"):
                text = "".join(node.text or "" for node in sibling.iter(qn("w:t"))).strip()
                if text.startswith("20. 核心实现代码"):
                    toc_paragraphs.append(sibling)
                    break
        if toc_paragraphs:
            new_toc = deepcopy(toc_paragraphs[-1])
            for node in new_toc.iter(qn("w:t")):
                node.text = "21. 测试证据与代码定位"
                break
            toc_paragraphs[-1].addnext(new_toc)

    target = next(p for p in doc.paragraphs if p.text.strip() == "核心实现代码")
    elements = []
    add_paragraph(doc, elements, page_break=True)
    add_paragraph(doc, elements, "DMBOK2 数据管理审核与整改", "Heading 1")
    add_paragraph(
        doc,
        elements,
        "本章依据用户提供的《DMBOK2.Rev 数据管理知识体系指南》所覆盖的数据治理、数据安全、数据集成与互操作、参考数据和主数据、元数据以及数据质量等知识领域，对当前可运行系统做工程审核。文中仅做框架概念映射和证据核对，不复制书中原文，也不将工程实现表述为 DAMA 认证、等保测评或法律合规结论。",
    )
    add_paragraph(doc, elements, "审核范围与总体结论", "Heading 2")
    add_paragraph(
        doc,
        elements,
        "整改后系统形成“治理责任—元数据登记—质量规则—问题发现—责任认领—处置证据—审核留痕—反馈回流”的可运行链路。当前纳入 14 项工程控制，其中 12 项达到 IMPLEMENTED 或 MONITORED；真实业务系统连接器和备份恢复演练保留为 DESIGNED，不用模拟结果替代现场验收。",
    )
    add_paragraph(doc, elements, "V5.1 统一 Agent 与可信数据流", "Heading 2")
    add_paragraph(
        doc,
        elements,
        "V5.1 将原先分散的智能检索、自然语言跨系统分发、治理编排和图文智能填报收敛为“主数据智能 Agent”统一入口。用户可输入文本、使用中文语音转写或上传工业铭牌/物料图片；系统先识别 SEARCH、DISTRIBUTE、GOVERN 等意图，再调用相应后端工具。检索可直接执行，分发与治理严格采用“生成计划—展示候选—人工选择—显式确认—执行留痕”的控制模式，避免大模型直接触发高风险业务操作。",
    )
    add_table(
        doc,
        elements,
        ["输入/任务", "处理链路", "输出证据", "失效保护"],
        [
            ("自然语言检索", "实体抽取 → 模糊/结构化检索 → 主数据排序 → standard_kb RAG", "黄金主数据候选、综合分、匹配理由、标准编号与版本", "在线模型失败时使用本地特征与规则检索"),
            ("跨系统分发", "意图识别 → 相似候选 → master×系统×工厂任务拆分 → 勾选确认", "任务载荷、执行状态、目标工厂、哈希凭证", "候选不自动执行；服务端重验权限、工厂范围和主数据状态"),
            ("治理编排", "目标识别 → 八步闭环计划 → 状态机推进", "步骤状态、责任主体、质量问题、审计事件", "越级状态迁移与伪造 actor/reviewer 被后端拒绝"),
            ("图文智能填报", "PaddleOCR/Qwen-VL → 字段抽取 → 标准化 → 联合检索或转申请", "原文、品牌、型号、压力、材质、口径、品类与置信度", "真实 OCR 不可用时明确标记规则降级，不伪装为模型识别"),
            ("中文语音输入", "Qwen ASR 或浏览器语音 → 文本草稿 → 用户确认", "可编辑的指令文本", "语音不会自动提交分发或治理操作"),
        ],
    )

    add_paragraph(doc, elements, "企业 RAG 与主数据检索协同", "Heading 2")
    add_paragraph(
        doc,
        elements,
        "知识库页面负责版本、索引和检索效果验证，面向业务人员的检索入口统一放在 Agent 工作台。/api/search 在一次请求中同时返回主数据候选和 knowledge_references；即使当前没有上传 CSV，用户仍可检索已发布的标准知识。标准依据包含 reference_id、standard_version、相关度及访问过滤标记，避免把向量分数当作无来源答案。",
    )

    add_paragraph(doc, elements, "关系血缘、桑基数据流与哈希审计", "Heading 2")
    add_paragraph(
        doc,
        elements,
        "治理血缘采用双视图而非单一图形：桑基视图面向管理汇报，展示多源接入、识别、Embedding 入库、图谱查重、审核、质量评估、跨厂分发与反馈回流的数量变化；关系视图面向运维排查，展示来源记录、黄金主数据、系统、工厂、治理步骤和反馈节点的网状关联。每个治理步骤关联哈希指纹和审计高度，可从聚合流量下钻到来源业务主键与处理证据。图中数量由当前批次动态聚合，不绑定 234 条样例数据。",
    )

    add_paragraph(doc, elements, "可视化展示层与合规呈现", "Heading 2")
    add_paragraph(
        doc,
        elements,
        "前端不设置独立的“治理与合规”业务页面，而是在接入、首页摘要、血缘审计中心和受保护接口中按任务呈现数据分级、处理目的、责任主体、AI 出域策略、来源追溯、保留期限、法律保全和审计链状态。该方式避免把控制条款与日常业务割裂，同时技术文档保留完整 DMBOK2 映射、残余风险和正式测评边界。",
    )

    add_paragraph(doc, elements, "运行与交付架构", "Heading 2")
    add_table(
        doc,
        elements,
        ["层级", "当前实现", "企业演进方向"],
        [
            ("体验与可视化", "单页 HTML、深浅主题、统一 Agent、桑基/关系双视图、治理报告", "前后端资源分离、严格 CSP、统一设计系统"),
            ("业务服务", "Flask API、八步状态机、审核、生命周期、分发、反馈", "多实例服务、任务队列、连接器编排与幂等重试"),
            ("AI 与知识", "通义 Embedding/Qwen Agent/Qwen ASR/Qwen-VL 可选，本地特征与规则降级", "模型网关、Prompt/模型版本治理、评测与成本配额"),
            ("数据与图", "SQLite 权威库、向量表、NetworkX 关系计算、版本化 standard_kb", "PostgreSQL/pgvector、图数据库、对象存储与 CDC"),
            ("安全与可信", "四角色会话、RBAC、工厂隔离、CSRF、HMAC 审计、保留与保全", "TLS/KMS、集中 IAM、SIEM、WORM 外部锚定、加密备份"),
            ("部署", "Windows start.bat、macOS start.command、Linux start.sh、Docker/Waitress", "反向代理、容器编排、监控告警和灾备演练"),
        ],
    )
    add_paragraph(doc, elements, "核心代码定位（按符号而非易漂移行号）", "Heading 2")
    add_table(
        doc,
        elements,
        ["能力", "实现符号", "验证入口"],
        [
            ("统一 Agent 路由与计划", "app.py: unified_agent_query、agent_plan、transcribe_agent_audio", "test_api.py: 统一 Agent 与单任务分发测试"),
            ("语义与降级", "app.py: SemanticEngine、api_semantic、_search_standard_kb", "test_api.py / test_chaos.py: 超时、畸形响应与降级"),
            ("OCR 三级适配", "app.py: OCREngine、api_ocr、api_ocr_install", "test_api.py / test_chaos.py: OCR 合同、Worker、安装权限"),
            ("版本化 RAG", "enterprise_rag.py: EnterpriseRAG.import/publish/rollback/search", "test_enterprise.py: RAG 版本与配置权限"),
            ("质量问题闭环", "enterprise_governance.py: capture_batch_issues、update_issue", "test_enterprise.py: DMBOK 控制目录与问题状态"),
            ("企业安全", "enterprise_security.py: EnterpriseSecurity、has_permission；app.py: attach_trace_context", "test_enterprise.py: 四账号、CSRF、RBAC、越权与篡改"),
            ("血缘与审计", "app.py: api_graph、api_lineage、_append_audit_block", "test_api.py / test_chaos.py: 图谱、来源追溯与哈希验证"),
        ],
    )

    add_table(
        doc,
        elements,
        ["DMBOK2 知识领域", "系统证据", "当前状态", "仍需企业现场完成"],
        [
            ("数据治理", "四角色职责、控制目录、责任人与证据、成熟度评估", "已实现", "数据治理委员会章程与业务负责人签字"),
            ("数据架构", "Flask 服务、SQLite 权威库、可视化展示层、接口契约", "已实现", "与企业总体数据架构及容量规划对齐"),
            ("数据存储与操作", "事务持久化、批次隔离、健康检查、容器部署", "部分实现", "加密备份、恢复演练、RPO/RTO 验收"),
            ("数据安全", "账户会话、RBAC、工厂范围、CSRF、分级、保留与法律保全", "已实现工程控制", "TLS、密钥托管、等保定级测评与法务确认"),
            ("数据集成与互操作", "动态字段映射、CSV 编码适配、意图分发、来源追溯", "演示闭环", "SAP/EAM/MES/WMS 真实连接器与 SLA"),
            ("参考数据与主数据", "SY/T 分类、黄金编码、审核、生命周期和跨厂分发", "已实现", "企业编码标准与主数据权威域审批"),
            ("元数据管理", "业务/技术/安全元数据目录、血缘、来源业务主键", "已实现", "与企业元数据平台交换及影响分析"),
            ("数据质量", "9 条规则、问题工单、责任人、截止时间、处置证据", "已实现", "业务真值集、阈值验收和持续 KPI 基线"),
            ("文档与内容生命周期", "RAG 来源、版本、哈希、发布、失效与回滚", "已实现", "标准文件版权、授权和归档制度确认"),
        ],
    )

    add_paragraph(doc, elements, "治理组织、责任与度量", "Heading 2")
    add_paragraph(
        doc,
        elements,
        "系统把职责绑定到登录身份，不再信任前端提交的 actor、reviewer 或 plant_code。每个控制包含责任角色、执行角色、KPI、目标值、成熟度和证据接口；工程覆盖率通过受保护接口和技术文档核验，该百分比只代表控制证据覆盖，不代表企业数据治理成熟度认证分数。",
    )
    add_table(
        doc,
        elements,
        ["角色", "数据范围", "主要职责", "关键限制"],
        [
            ("集团管理员 GROUP_ADMIN", "集团及授权工厂", "账号、策略、控制评估、知识版本发布", "高风险操作需 CSRF、签名审计和保留策略约束"),
            ("集团审批人 GROUP_APPROVER", "集团及授权工厂", "主数据审核、生命周期审批、标准治理", "不能伪造其他审批人或越过状态机"),
            ("工厂数据专员 PLANT_STEWARD", "所属工厂", "数据接入、问题认领、修复、反馈回流", "请求 GROUP 或其他工厂时后端返回 403"),
            ("审计员 AUDITOR", "授权审计范围", "只读审计、控制证据、链完整性验证", "不能修改质量问题、主数据或控制状态"),
        ],
    )

    add_paragraph(doc, elements, "业务术语、元数据与血缘", "Heading 2")
    add_paragraph(
        doc,
        elements,
        "元数据目录登记 material_code、material_name、description、brand、model、category、system_source、plant_code、source_record_id 和 source_url 等关键元素，记录业务定义、技术类型、责任角色、安全级别和质量规则。问题工单保留 batch_id、record_id、来源系统、来源业务主键和工厂范围，可从黄金主数据反向追溯到 ERP/EAM/采购源记录；血缘图和治理报告共用同一批次证据。",
    )

    add_paragraph(doc, elements, "数据质量规则与问题闭环", "Heading 2")
    add_paragraph(
        doc,
        elements,
        "系统启用 9 条质量规则，覆盖物料编码和名称必填、来源系统可追溯、工厂范围、分类标准化、描述完整性、品牌型号完整性、创建时间格式以及来源编码唯一性。接入和治理时按当前数据集动态生成问题，不把 234 条测试数据写死。问题状态为 OPEN → ACKNOWLEDGED → RESOLVED/WAIVED，记录责任人、严重度、截止时间、处置说明和证据引用，并写入签名安全审计。",
    )
    add_paragraph(
        doc,
        elements,
        "质量报告中的“准确性代理”仅衡量结构异常和字段规则，不等于模型 Precision/Recall。正式业务准确率须由独立抽样、双人标注、争议仲裁和业务签字形成验证集，再按样本时间范围发布。",
    )

    add_paragraph(doc, elements, "数据安全、保留与可信审计", "Heading 2")
    add_paragraph(
        doc,
        elements,
        "企业模式提供四个独立账号、密码哈希、会话撤销、密码变更、登录锁定、后端 RBAC、工厂范围交集、CSRF/Origin 校验和安全事件记录。数据按 PUBLIC、INTERNAL、CONFIDENTIAL、RESTRICTED 分级；外部 AI 调用先执行分级与最小化策略，受限数据默认本地处理。保留策略支持 dry-run、到期处置和 legal hold，法律保全优先于清理。审计使用 HMAC 签名链并覆盖身份、授权拒绝、治理、审核、分发、反馈和保留操作。生产部署仍需 TLS、密钥托管、加密备份、日志外送和独立 WORM 锚定。",
    )

    add_paragraph(doc, elements, "企业 RAG 标准知识库配置", "Heading 2")
    add_paragraph(
        doc,
        elements,
        "RAG 知识库不再由启动时的 10 条常量覆盖。标准文件进入隔离版本，经过 DRAFT → VALIDATED → PUBLISHED 后才参与正式分类；发布可回滚，旧版本保留审计证据。检索先做工厂与密级 ACL，再进行规则、词法和向量三路融合，返回标准号、条款/分类编码、来源行号、版本和内容哈希引用。",
    )
    add_table(
        doc,
        elements,
        ["配置项", "默认值", "用途"],
        [
            ("MDM_RAG_SOURCE_PATH / MDM_RAG_VERSION", "SY/T 分类树 / 2018", "启动标准源与版本标签"),
            ("MDM_RAG_DIMENSION", "384", "本地特征向量维度"),
            ("MDM_RAG_RULE_WEIGHT", "0.45", "规则命中权重"),
            ("MDM_RAG_LEXICAL_WEIGHT", "0.30", "词法与别名权重"),
            ("MDM_RAG_VECTOR_WEIGHT", "0.25", "向量语义权重"),
            ("MDM_RAG_TOP_K", "3", "返回候选数量"),
            ("MDM_RAG_AUTO_ACCEPT_THRESHOLD", "0.68", "自动接受最低综合分"),
            ("MDM_RAG_MINIMUM_MARGIN", "0.08", "Top1 与 Top2 最小差距"),
            ("MDM_RAG_ALLOWED_PLANTS", "*", "知识版本允许的工厂范围"),
            ("MDM_RAG_ALLOWED_CLASSIFICATIONS", "INTERNAL,CONFIDENTIAL,RESTRICTED", "调用者可检索的密级范围"),
            ("DASHSCOPE_API_KEY", "未配置", "可选通义向量；失败自动降级本地特征"),
        ],
    )

    add_paragraph(doc, elements, "新增接口与证据路径", "Heading 2")
    add_table(
        doc,
        elements,
        ["接口", "用途", "权限/证据"],
        [
            ("GET /api/governance/catalog", "控制、RACI、元数据和质量规则目录", "data.read；支持当前批次问题汇总"),
            ("GET /api/governance/issues", "按批次、工厂、状态查询质量问题", "服务端按登录主体收敛范围"),
            ("PATCH /api/governance/issues/<id>", "认领、解决或豁免问题", "review.decide；签名事件留痕"),
            ("PATCH /api/governance/controls/<code>", "评估控制成熟度和证据", "compliance.manage；仅授权角色"),
            ("GET /api/compliance/status", "数据分级、AI 出域、保留、法律保全和 DMBOK 摘要", "受保护读取"),
            ("/api/knowledge/*", "标准版本、配置、导入、验证、发布、回滚与索引管理", "knowledge.manage / data.read；版本化审计"),
        ],
    )

    add_paragraph(doc, elements, "整改验证与残余事项", "Heading 2")
    add_paragraph(
        doc,
        elements,
        "2026-08-08 自动化验证结果为 40/40：test_api.py 21 项、test_chaos.py 9 项、test_enterprise.py 10 项。前端保留接入、总览、知识库、Agent 工作台、审核、数据质量、生命周期和分发等业务导航，不设置独立的治理与合规页面；治理证据保留在后端接口、签名审计与本文档中。深色与浅色业务页面已完成桌面端可视检查；实际浏览器验证确认 Agent 可同时返回 18 条主数据候选和 3 条带版本/引用编号的 RAG 标准依据，控制台无错误。正式移动终端兼容仍应纳入交付验收。",
    )
    add_table(
        doc,
        elements,
        ["残余事项", "当前边界", "建议验收证据"],
        [
            ("真实业务连接器", "当前 SAP/EAM/MES/WMS 为可替换演示适配器", "接口联调、幂等、重试、对账和 SLA 报告"),
            ("业务准确率与 ROI", "演示指标不替代生产收益", "业务签字真值集、岗位计时、成本表和三档 ROI"),
            ("备份与灾备", "具备持久化但未完成企业恢复演练", "加密备份、恢复截图、RPO/RTO 记录"),
            ("正式安全合规", "已形成工程控制映射，不宣称通过认证", "单位定级、等保测评、法务和数据出境评估"),
            ("规模化向量检索", "SQLite 精确扫描适合当前演示规模", "大规模时迁移 pgvector/HNSW 并提供基准测试"),
        ],
    )
    add_paragraph(doc, elements, page_break=True)

    for element in elements:
        target._p.addprevious(element)

    add_paragraph(
        doc,
        [],
        "本节只列可重复执行的验证入口；核心代码定位已在第 19 章按类名和函数名给出，避免以易漂移的行号作为证据。执行测试前使用 backend/.env.example 创建本地配置，并确保真实密钥只保存在已忽略的 backend/.env。",
    )
    verification_table = doc.add_table(rows=1, cols=3)
    for index, header in enumerate(("验证项", "命令或接口", "当前结果")):
        verification_table.rows[0].cells[index].text = header
    verification_rows = (
        ("核心 API 回归", "python test_api.py", "21/21 OK"),
        ("故障与降级", "python test_chaos.py", "9/9 OK"),
        ("企业安全与治理", "python test_enterprise.py", "10/10 OK"),
        ("Python 语法", "python -m py_compile backend/app.py backend/enterprise_*.py", "通过"),
        ("前端脚本", "Node.js 解析 index.html 内联脚本", "通过"),
        ("运行健康", "GET /api/health", "ready=true；version=5.1；security_mode=enterprise"),
        ("浏览器可视检查", "浅色主题 Agent 工作台与 RAG 联合检索", "无重叠；控制台无错误"),
    )
    for values in verification_rows:
        cells = verification_table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    style_table(verification_table)

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    doc.save(str(TARGET))
    print(TARGET)


if __name__ == "__main__":
    build_document()
